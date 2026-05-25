from docx import Document


class WordParser:

    def parse(self, file_path):

        doc = Document(file_path)

        text = "\n".join(

            [p.text for p in doc.paragraphs]
        )

        return [
            {
                "text": text
            }
        ]



# from docx import Document

# def parse_docx(file_path):

#     doc = Document(file_path)

#     text = "\n".join(
#         [p.text for p in doc.paragraphs]
#     )

#     return text